import requests
from bs4 import BeautifulSoup
import csv
from fuzzywuzzy import fuzz
from docx import Document
import glob
from difflib import SequenceMatcher


def similar(a, b):
    return fuzz.token_sort_ratio(a, b)

def save_csv_deputados():
    url_to_scrape = 'http://www.asamblea.go.cr/ca/lists/diputados/diputadas%20y%20diputados.aspx'

    r = requests.get(url_to_scrape)

    soup = BeautifulSoup(r.text, "html.parser")
    #search table
    iframexx = soup.find_all('table')

    diputados = []
    dip_names = []
    diputados.append([["email"],["nombre"],["provincia"],["partido"]])
    # the second table corresponds to the one with the senators
    for row in iframexx[2].find_all('tr')[1:]:
        listrow = row.find_all('td')
        # if the row has more than 2 columns, it is probably alright

        if (len(listrow)>2):
            demail = listrow[1].get_text()
            dname  = listrow[2].get_text()
            dplace = listrow[3].get_text()
            dparty = listrow[4].get_text()
            # some have a web site on the 5th column, but that's alright
            diputados.append([demail, dname, dplace, dparty])
            dip_names.append(dname)
    print(diputados)
    with open('diputados.csv', 'w') as myFile:
        writer = csv.writer(myFile)
        writer.writerows(diputados)
    return dip_names









if __name__ == '__main__':
    dip_names = save_csv_deputados()
    print(dip_names)
    path = "*PLENARIO*"
    fname_list = glob.glob(path)
    with open('ausencias.csv', 'w') as myFile:
        writer = csv.writer(myFile)

        dip_names.append("PLENARY")
        print(dip_names)
        writer.writerow(dip_names)
        fname_list.sort()
        for fname in fname_list:
            document = Document(fname)
            document.save(fname)
            empty_list = []
            empty_list = [0] * len(dip_names)
            for row in document.tables[0].rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        #print(para.text)
                        diputado_index = 0
                        for diputado in dip_names:
                            if (similar(para.text.lower(), diputado.lower()))>90:
                                # print("")
                                # print("@@@@@@@@@@@@@@@@@@@@@")
                                # print(diputado.lower())
                                # print(para.text.lower())
                                # print("@@@@@@@@@@@@@@@@@@@@@")
                                # print(similar(para.text, diputado.lower()))
                                # print("")
                                empty_list[diputado_index]=1
                        # print(similar(para.text, diputado.lower()))
                            #fuzz.ratio("this is a test", "this is a test!")
                            diputado_index = diputado_index + 1

            empty_list.append(fname)
            print(empty_list)
            writer.writerow(empty_list)
            diputado_index = 0
            # for falta in empty_list:
            #     if falta == 0:
            #         print(dip_names[diputado_index])
            #     diputado_index = diputado_index + 1
