from docx import Document
import os
import glob



path = "*PLENARIO*"
for fname in glob.glob(path):
    print(fname)

document = Document('2017-2018-PLENARIO-SESION-162.docx')
document.save('2017-2018-PLENARIO-SESION-162.docx')

for row in document.tables[0].rows:
  for cell in row.cells:
    for para in cell.paragraphs:
      print(para.text)